import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Inches
from datetime import datetime as dt
import os
import shutil
from datetime import datetime
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt


n = 0
for yr in os.listdir("facturas"):
    if os.path.isdir(f"facturas/{yr}"):
        for mt in os.listdir(f"facturas/{yr}"):
            for day in os.listdir(f"facturas/{yr}/{mt}"):
                if os.path.isdir(f"facturas/{yr}/{mt}/{day}"):
                    for z in os.listdir(f"facturas/{yr}/{mt}/{day}"):
                        if z.endswith("xlsx"):
                            n += 1
                else:
                    continue


invoice_path = os.listdir("facturas")[-1]

yr = os.path.splitext(invoice_path)[0].split("-")[2]
mt = "/".join(os.path.splitext(invoice_path)[0].split("-")[:2])

df = pd.read_excel(os.path.abspath(f"facturas/{invoice_path}"), sheet_name="factura").ffill().set_index("DIRECCION CLIENTE")
clients = pd.read_excel(os.path.abspath("inventario.xlsm"), sheet_name="clientes")
inventario = pd.read_excel(os.path.abspath("inventario.xlsm"), sheet_name="inventario")


address = df.index.values[0] if df.index.values[0] != np.nan != 0 else "Sin Direccion"

nombre = df["NOMBRE CLIENTE"].iloc[0] if df["NOMBRE CLIENTE"].iloc[0] != np.nan else "Sin nombre"


if address != "Sin Direccion":
    NIT_cliente = clients.loc[clients["DIRECCION"] == address, "NIT"].iloc[0]
elif nombre != "Sin nombre":
    NIT_cliente = clients.loc[clients["NOMBRE CLIENTE"] == nombre, "NIT"].iloc[0]

if address != "Sin Direccion":
    tel = clients.loc[clients["DIRECCION"] == address, "TELEFONO"].iloc[0] if not pd.isna(clients.loc[clients["DIRECCION"] == address, "TELEFONO"].iloc[0]) else "Sin Telefono"
    
elif nombre != "Sin nombre":
    tel = clients.loc[clients["NOMBRE CLIENTE"] == nombre, "TELEFONO"].iloc[0] if not pd.isna(clients.loc[clients["NOMBRE CLIENTE"] == nombre, "TELEFONO"].iloc[0]) else "Sin Telefono"
    



df.drop("NOMBRE CLIENTE", axis=1, inplace=True)

def format_number(x):
    return f"${x:,.0f}"

df[["PRECIO UNIDAD", "TOTAL"]] = df[["PRECIO UNIDAD", "TOTAL"]].map(format_number)

invoice = Document()

section = invoice.sections[0]

section.left_margin = Inches(0.2)
section.right_margin = Inches(8.0)
section.top_margin = Inches(0.2)
# section.bottom_margin = Inches(1.0)

NIT = "43.824.404-5"
direccion = "Carrera 53 #59-31, La Candelaria"
celular = "3113404920"


table = invoice.add_table(rows=1, cols=1)
table.autofit = False


for i in table.columns[0].cells:
    i.width = Inches(3.5)

paragraph_pic = table.cell(0,0).paragraphs[0]

run = paragraph_pic.add_run()

run.font.name = "Corbel (Headings)"

run.add_picture("logo.png", height=Inches(1.03), width=Inches(1.33))

paragraph_pic.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


paragraph_info = table.cell(0,0).add_paragraph()

run = paragraph_info.add_run("Distribuidora Antaño-Moore")
run.font.size = Pt(8)
run.font.name = "Corbel (Body)"
run.add_break()
run.add_text("NIT" + NIT)
run.add_break()
run.add_text(direccion)
run.add_break()
run.add_text("Cel: " + celular)
run.add_break()
run.add_text("Factura No " + str(n+1))
run.add_break()
run.add_text("__________________________________________________________")

paragraph_info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

paragraph_des = table.cell(0,0).add_paragraph()

run = paragraph_des.add_run("Fecha: " + datetime.now().strftime("%d/%M/%Y"))
run.font.size = Pt(10)
run.font.name = "Corbel (Body)"
run.add_break()
run.add_break()

run_2 = paragraph_des.add_run("CLIENTE:")
run_2.font.bold = True
run_2.font.size = Pt(10)
run_2.font.name = "Corbel (Body)"
run_2.add_break()

run_3 = paragraph_des.add_run(f"Dirección: {address}")
run_3.add_break()
run_3.add_text(F"NIT: {NIT_cliente}")
run_3.add_break()
run_3.add_text(f"Teléfono: {tel}")
run_3.font.size = Pt(10)
run_3.font.name = "Corbel (Body)"


invoice.add_paragraph("\n")

table_2 = invoice.add_table(rows=1, cols=4)
table_2.autofit = False



for i,j in zip(table_2.rows[0].cells, df[["CANT", "DESCRIPCIÓN", "PRECIO UNIDAD", "TOTAL"]].columns):
    i.text = j


widths = [Inches(0.5), Inches(1.18), Inches(0.59), Inches(0.65)] 
for i, width in enumerate(widths):
    for cell in table_2.columns[i].cells:
        cell.width = width


for i in df[["CANT", "DESCRIPCIÓN", "PRECIO UNIDAD", "TOTAL"]].values:
    new_row = table_2.add_row().cells
    for j,k in zip(new_row, i):
        j.text = f"{k}"


for row in table_2.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(8)
                run.font.name = "Corbel (Body)"
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

table_t = invoice.add_table(rows=1, cols=4)
table_t.autofit = False

widths = [Inches(0), Inches(0.59+1.18), Inches(0), Inches(0.49)]
for i, width in enumerate(widths):
    for cell in table_t.columns[i].cells:
        cell.width = width

table_t.rows[0].cells[-3].text = "TOTAL A PAGAR"

table_t.rows[0].cells[-1].text = f"{format_number(df['TOTAL'].str.replace(',', '.').str[1:].astype(float).sum())}"

for row in table_t.rows:
    for cell in row.cells[0:-1]:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(8)
                run.font.name = "Corbel (Body)"
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                run.font.bold = True

for row in table_t.rows:
    for cell in row.cells[1:]:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(8)
                run.font.name = "Corbel (Body)"
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

if os.path.exists(f"facturas/{yr}/{mt}"):
    pass
else:
    os.makedirs(f"facturas/{yr}/{mt}")

invoice.save(f"facturas/{yr}/{mt}/{os.path.splitext(invoice_path)[0]}_{n+1}.docx")

name_with_number = f"{os.path.splitext(invoice_path)[0]}_{n+1}.xlsx"

os.rename(f"facturas/{invoice_path}", f"facturas/{name_with_number}")

shutil.move(f"facturas/{name_with_number}", f"facturas/{yr}/{mt}/")