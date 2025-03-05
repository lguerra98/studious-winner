import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Inches
from datetime import datetime as dt
import os
import shutil
from datetime import datetime
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt

def generar_factura():


    path = [i for i in os.listdir() if i.startswith(pd.Timestamp.now().strftime("%Y%m%d"))][0]
    headers = ["cant", "producto", "valor unidad", "unidad"]

    df = pd.read_csv(path, header=None, delimiter=";")
    df_clientes = pd.read_excel("inventario.xlsm", sheet_name="clientes")

    df.columns = headers
    df["total"] = df["cant"]*df["valor unidad"]

    df["cant"] = df["cant"].astype(str) + " " + df["unidad"]
    df.drop("unidad", axis=1, inplace=True)

    nit = int(path.split("_")[1])
    uid = path.split("_")[0]

    def format_number(x):
        return f"${x:,.0f}"
        
    invoice = Document()

    section = invoice.sections[0]

    section.left_margin = Inches(0.2)
    section.right_margin = Inches(8.0)
    section.top_margin = Inches(0.2)
    # section.bottom_margin = Inches(1.0)

    NIT = "43.824.404-5"
    direccion = "Carrera 53 #59-31, La Candelaria"
    celular = "3113404920"


    table = invoice.add_table(rows=1, cols=1)
    table.autofit = False


    for i in table.columns[0].cells:
        i.width = Inches(3.5)

    df[["valor unidad", "total"]] = df[["valor unidad", "total"]].map(format_number)


    paragraph_pic = table.cell(0,0).paragraphs[0]

    run = paragraph_pic.add_run()

    run.font.name = "Cambria"

    run.add_picture("logo.png", height=Inches(1.03), width=Inches(1.33))

    paragraph_pic.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


    paragraph_info = table.cell(0,0).add_paragraph()

    run = paragraph_info.add_run("Distribuidora Antaño-Moore")
    run.font.size = Pt(8)
    run.font.name = "Cambria"
    run.add_break()
    run.add_text(f"NIT {NIT}")
    run.add_break()
    run.add_text(direccion)
    run.add_break()
    run.add_text(f"Cel: {celular}")
    run.add_break()
    run.add_text(f"Factura No {uid}")
    run.add_break()
    run.add_text("__________________________________________________________")

    paragraph_info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph_des = table.cell(0,0).add_paragraph()

    run = paragraph_des.add_run("Fecha: " + datetime.now().strftime("%d/%M/%Y"))
    run.font.size = Pt(10)
    run.font.name = "Cambria"
    run.add_break()
    run.add_break()

    run_2 = paragraph_des.add_run("CLIENTE:")
    run_2.font.bold = True
    run_2.font.size = Pt(10)
    run_2.font.name = "Cambria"
    run_2.add_break()

    run_3 = paragraph_des.add_run(f"Dirección: {df_clientes.query(f"NIT == {nit}")["DIRECCION"].iloc[0]}")
    run_3.add_break()
    run_3.add_text(F"NIT: {nit}")
    run_3.add_break()
    run_3.add_text(f"Teléfono: {df_clientes.query(f"NIT == {nit}")["TELEFONO"].iloc[0]}")
    run_3.font.size = Pt(10)
    run_3.font.name = "Cambria"


    invoice.add_paragraph("\n")

    table_2 = invoice.add_table(rows=1, cols=4)
    table_2.autofit = False

    for i,j in zip(table_2.rows[0].cells, df.columns):
        i.text = j.upper()

    widths = [Inches(0.7), Inches(1.18), Inches(0.59), Inches(0.65)] 

    for i, width in enumerate(widths):
        for cell in table_2.columns[i].cells:
            cell.width = width


    for i in df.values:
        new_row = table_2.add_row().cells
        for j,k in zip(new_row, i):
            j.text = f"{k}"


    for row in table_2.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
                    run.font.name = "Cambria"
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    table_t = invoice.add_table(rows=1, cols=4)
    table_t.autofit = False

    widths = [Inches(0), Inches(0.59+1.18), Inches(0), Inches(0.49)]

    for i, width in enumerate(widths):
        for cell in table_t.columns[i].cells:
            cell.width = width

    table_t.rows[0].cells[-3].text = "TOTAL A PAGAR"

    table_t.rows[0].cells[-1].text = f"{format_number(df['total'].str.replace(',', '').str[1:].astype(float).sum())}"

    for row in table_t.rows:
        for cell in row.cells[0:-1]:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
                    run.font.name = "Cambria"
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    run.font.bold = True

    for row in table_t.rows:
        for cell in row.cells[1:]:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
                    run.font.name = "Cambria"
                    if run.text.startswith("T"):
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    else:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    date = pd.Timestamp.now().strftime("%Y-%m")

    if os.path.exists(f"facturas/{date}"):
        invoice.save(f"facturas/{date}/{uid}.docx")
    else:
        os.makedirs(f"facturas/{date}")
        invoice.save(f"facturas/{date}/{uid}.docx")
    
    os.remove(path=path)


if __name__ == "__main__":
    generar_factura()